import hashlib
import math
import re
from collections import Counter
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from flask import current_app
from sqlalchemy import delete

from app.ai.duplicates import OllamaEmbeddingProvider
from app.ai.ocr import OCR_MIME_TYPES, ocr_provider
from app.extensions import db
from app.models import (
    AuditLog,
    NotificationType,
    OutboxEvent,
    RagChunk,
    RagDocumentVersion,
    RagIngestionStatus,
)
from app.notifications.service import notify_user
from app.rag.storage import rag_document_path

RAG_INGESTION_EVENT = "IngestaoDocumentoRag"


class RagIngestionError(RuntimeError):
    pass


class NonRetryableRagError(RagIngestionError):
    pass


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    pages: list[dict]


def enqueue_ingestion(version: RagDocumentVersion) -> None:
    db.session.add(
        OutboxEvent(
            tenant_id=version.tenant_id,
            event_type=RAG_INGESTION_EVENT,
            aggregate_type="DocumentoRag",
            aggregate_id=str(version.id),
            payload={"versionId": str(version.id)},
        )
    )


def requeue_ingestion(version: RagDocumentVersion) -> None:
    version.ingestion_status = RagIngestionStatus.PENDENTE
    version.error = None
    version.started_at = None
    version.indexed_at = None
    enqueue_ingestion(version)


def execute_ingestion(version: RagDocumentVersion) -> None:
    if version.ingestion_status == RagIngestionStatus.INDEXADO:
        return
    version.ingestion_status = RagIngestionStatus.PROCESSANDO
    version.started_at = datetime.now(UTC)
    version.error = None
    db.session.flush()

    extracted = extract_document(rag_document_path(version.storage_key), version.mime_type)
    if len(extracted.text.strip()) < current_app.config["RAG_MIN_TEXT_CHARS"]:
        raise NonRetryableRagError("O documento não possui texto suficiente para indexação.")
    chunks = split_chunks(
        extracted.pages,
        current_app.config["RAG_CHUNK_SIZE_CHARS"],
        current_app.config["RAG_CHUNK_OVERLAP_CHARS"],
    )
    if not chunks:
        raise NonRetryableRagError("Não foi possível gerar fragmentos do documento.")
    provider = rag_embedding_provider()
    vectors = provider.embeddings([item["content"] for item in chunks])
    if len(vectors) != len(chunks):
        raise RagIngestionError("O provedor retornou embeddings incompletos.")

    db.session.execute(delete(RagChunk).where(RagChunk.version_id == version.id))
    for item, vector in zip(chunks, vectors, strict=True):
        db.session.add(
            RagChunk(
                tenant_id=version.tenant_id,
                version_id=version.id,
                position=item["position"],
                content=item["content"],
                content_checksum=hashlib.sha256(item["content"].encode("utf-8")).hexdigest(),
                page_start=item["pageStart"],
                page_end=item["pageEnd"],
                section=item["section"],
                embedding=vector,
                embedding_model=provider.model,
            )
        )
    version.extracted_text = extracted.text
    version.page_count = len(extracted.pages)
    version.embedding_model = provider.model
    version.chunk_count = len(chunks)
    version.ingestion_status = RagIngestionStatus.INDEXADO
    version.indexed_at = datetime.now(UTC)
    details = {
        "documentoId": str(version.document_id),
        "versaoId": str(version.id),
        "checksum": version.checksum,
        "modelo": provider.model,
        "paginas": version.page_count,
        "fragmentos": version.chunk_count,
    }
    db.session.add(
        AuditLog(
            tenant_id=version.tenant_id,
            user_id=version.created_by_id,
            action="rag_document.indexed",
            entity_type="rag_document_version",
            entity_id=str(version.id),
            after=details,
        )
    )
    notify_user(
        version.tenant_id,
        version.created_by_id,
        NotificationType.SISTEMA,
        "Documento indexado",
        f"{version.original_name} está disponível na base documental.",
        "rag_document_version",
        version.id,
    )


def fail_ingestion(version: RagDocumentVersion, error_message: str) -> None:
    version.ingestion_status = RagIngestionStatus.FALHOU
    version.error = error_message[:2000]


def extract_document(path: Path, mime_type: str) -> ExtractedDocument:
    if mime_type == "text/plain":
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError as error:
            raise NonRetryableRagError("O arquivo de texto deve estar em UTF-8.") from error
        return ExtractedDocument(text=text, pages=[{"pagina": 1, "texto": text}])
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            paragraphs = [
                item.text.strip() for item in Document(path).paragraphs if item.text.strip()
            ]
        except Exception as error:
            raise NonRetryableRagError("Não foi possível ler o documento DOCX.") from error
        text = "\n\n".join(paragraphs)
        return ExtractedDocument(text=text, pages=[{"pagina": 1, "texto": text}])
    if mime_type in OCR_MIME_TYPES:
        result = ocr_provider().extract(path, mime_type)
        return ExtractedDocument(text=result.text, pages=result.pages)
    raise NonRetryableRagError("Tipo de documento não compatível com a ingestão.")


def split_chunks(pages: list[dict], size: int, overlap: int) -> list[dict]:
    chunks = []
    position = 0
    for page in pages:
        text = re.sub(r"[ \t]+", " ", str(page.get("texto") or "")).strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(start + size, len(text))
            if end < len(text):
                boundary = text.rfind(" ", start + size // 2, end)
                if boundary > start:
                    end = boundary
            content = text[start:end].strip()
            if content:
                chunks.append(
                    {
                        "position": position,
                        "content": content,
                        "pageStart": int(page.get("pagina") or 1),
                        "pageEnd": int(page.get("pagina") or 1),
                        "section": page.get("secao"),
                    }
                )
                position += 1
            if end >= len(text):
                break
            start = max(end - overlap, start + 1)
    return chunks


class LocalHashEmbeddingProvider:
    model = "gabflow-hash-embedding-v1"

    def embeddings(self, texts: list[str]) -> list[list[float]]:
        return [_hash_embedding(text) for text in texts]


def rag_embedding_provider():
    if current_app.config["RAG_EMBEDDING_PROVIDER"].lower() == "local":
        return LocalHashEmbeddingProvider()
    return OllamaEmbeddingProvider(
        current_app.config["OLLAMA_BASE_URL"],
        current_app.config["AI_EMBEDDING_MODEL"],
        current_app.config["RAG_INGESTION_TIMEOUT_SECONDS"],
    )


def _hash_embedding(text: str, dimensions: int = 128) -> list[float]:
    tokens = re.findall(r"[a-zA-ZÀ-ÿ0-9]+", text.lower())
    counts = Counter(tokens)
    vector = [0.0] * dimensions
    for token, count in counts.items():
        digest = hashlib.sha256(token.encode("utf-8")).digest()
        index = int.from_bytes(digest[:4], "big") % dimensions
        vector[index] += float(count)
    norm = math.sqrt(sum(value * value for value in vector)) or 1.0
    return [round(value / norm, 8) for value in vector]
