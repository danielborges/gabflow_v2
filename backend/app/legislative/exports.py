from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models import LegislativeDraft


def draft_docx(draft: LegislativeDraft) -> BytesIO:
    document = Document()
    document.add_heading(draft.title, level=1)
    document.add_paragraph(draft.document_type.value.replace("_", " ").title())
    for block in (draft.content or "").split("\n"):
        document.add_paragraph(block)
    if draft.justification:
        document.add_heading("Justificativa", level=2)
        document.add_paragraph(draft.justification)
    if draft.legal_basis:
        document.add_heading("Fundamentação informada", level=2)
        for source in draft.legal_basis:
            document.add_paragraph(
                f"{source.get('titulo', 'Fonte')} — {source.get('referencia') or 'sem referência'}",
                style="List Bullet",
            )
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream


def draft_pdf(draft: LegislativeDraft) -> BytesIO:
    stream = BytesIO()
    styles = getSampleStyleSheet()
    story = [Paragraph(_escape(draft.title), styles["Title"]), Spacer(1, 0.4 * cm)]
    for block in (draft.content or "").split("\n"):
        if block.strip():
            story.extend([Paragraph(_escape(block), styles["BodyText"]), Spacer(1, 0.2 * cm)])
    if draft.justification:
        story.extend(
            [
                Spacer(1, 0.4 * cm),
                Paragraph("Justificativa", styles["Heading2"]),
                Paragraph(_escape(draft.justification), styles["BodyText"]),
            ]
        )
    SimpleDocTemplate(
        stream,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=draft.title,
    ).build(story)
    stream.seek(0)
    return stream


def _escape(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
